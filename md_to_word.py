import argparse
from pathlib import Path
from markdown import markdown
from docx import Document
from bs4 import BeautifulSoup


def markdown_to_docx(md_file_path):
    md_path = Path(md_file_path)
    if not md_path.exists():
        print(f"Error: File '{md_file_path}' does not exist.")
        return

    with md_path.open('r', encoding='utf-8') as md_file:
        md_content = md_file.read()

    # Convert markdown to HTML
    html_content = markdown(md_content)

    # Parse HTML content
    soup = BeautifulSoup(html_content, 'html.parser')

    # Create a new Word document
    doc = Document()

    # Add paragraphs to the Word document from HTML
    for element in soup.descendants:
        if element.name == 'p':
            doc.add_paragraph(element.get_text())
        elif element.name == 'h1':
            doc.add_heading(element.get_text(), level=1)
        elif element.name == 'h2':
            doc.add_heading(element.get_text(), level=2)
        elif element.name == 'h3':
            doc.add_heading(element.get_text(), level=3)
        elif element.name == 'li':
            doc.add_paragraph(element.get_text(), style='List Bullet')

    # Save the Word document in the same directory with .docx extension
    docx_path = md_path.with_suffix('.docx')
    doc.save(docx_path)
    print(f"Markdown file '{md_file_path}' has been converted to '{docx_path}'.")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description='Convert Markdown file to DOCX.')
    parser.add_argument('markdown_file', help='Path to the Markdown file to convert.')

    args = parser.parse_args()

    markdown_to_docx(args.markdown_file)

